import os
import datetime
import glob
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent

def create_word_report():
    today = datetime.datetime.now().strftime("%Y%m%d")
    md_file_name = f"chatbot_defect_report_{today}.md"
    md_file_path = BASE_DIR / md_file_name
    
    if not md_file_path.exists():
        print(f"❌ 금일 결함 보고서 마크다운 파일이 존재하지 않습니다: {md_file_name}")
        return

    # 1. 파일 전체를 읽어 요약 정보(결함 건수, 원인 등)를 먼저 수집합니다.
    with open(md_file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    defect_count = 0
    reasons = set()
    for line in lines:
        if line.startswith("## 결함 기록:"):
            defect_count += 1
        elif "| **원인 추정** |" in line:
            parts = line.split("|")
            if len(parts) >= 3:
                reason = parts[2].strip()
                if reason:
                    reasons.add(reason)

    doc = Document()
    
    # [문서 제목]
    title = doc.add_heading(f"챗봇 답변 품질 일일 결함 보고서 ({today})", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # [결함 요약 (Summary)]
    doc.add_heading("1. 결함 요약 (Summary)", level=1)
    doc.add_paragraph(f"• 금일 발생한 총 품질 이상(결함) 건수: {defect_count}건")
    if reasons:
        doc.add_paragraph("• 주요 원인 요약:")
        for r in list(reasons)[:3]:  # 최대 3개까지만 요약
            doc.add_paragraph(f"  - {r}")
    
    doc.add_paragraph()
    doc.add_heading("2. 결함 상세 내역", level=1)

    current_table = None
    extracted_reason = ""
    extracted_action = ""
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith("## 결함 기록:"):
            doc.add_heading(line.replace("## ", ""), level=2)
            current_table = None
            extracted_reason = "원인 파악 중"
            extracted_action = "추가 검증 필요"
        elif line.startswith("# ") or line.startswith("본 문서는"):
            pass # 스킵 (제목 및 마크다운 기본 설명)
        elif line.startswith("|"):
            if "---" in line:
                continue # 구분선 스킵
                
            cells = [cell.strip() for cell in line.split("|")[1:-1]]
            
            # 원인 추정과 조치 방안 내용을 캡처하여 나중에 QA 코멘트로 사용
            if len(cells) >= 2:
                header = cells[0].replace("**", "").strip()
                content = cells[1].replace("**", "").strip()
                if "원인 추정" in header:
                    extracted_reason = content
                elif "조치 방안" in header:
                    extracted_action = content

            if not current_table:
                # 새 표 생성
                current_table = doc.add_table(rows=1, cols=len(cells))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    hdr_cells[i].text = cell_text.replace("**", "")
            else:
                row_cells = current_table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        row_cells[i].text = cell_text.replace("**", "")
        elif line == "---":
            # 한 결함 내역 파싱이 끝남 (표 아래에 QA 코멘트 자동 생성)
            current_table = None
            doc.add_paragraph()
            p = doc.add_paragraph()
            qa_text = f"💡 QA 엔지니어 검토 의견: 본 품질 이상은 [{extracted_reason}] 문제로 기인한 것으로 분석됩니다. 따라서 [{extracted_action}] 조치를 담당 개발 파트에 강력히 권고하며, 해당 조치 완료 즉시 유사 변형 질문 50건에 대한 추가 재검증(Regression Test)을 수행할 예정입니다."
            run = p.add_run(qa_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0, 0, 255) # 파란색 글씨로 강조
            doc.add_paragraph() # 여백
        else:
            if current_table is None:
                pass # 일반 텍스트는 보통 테이블 파싱 후에는 무시하거나 필요시 추가

    word_file_name = f"chatbot_defect_report_{today}.docx"
    word_file_path = BASE_DIR / word_file_name
    doc.save(word_file_path)
    print(f"Word 보고서가 성공적으로 생성되었습니다: {word_file_name}")

if __name__ == "__main__":
    create_word_report()
